"""Извлечение текста из загруженного учебного плана.

Поддержка: .txt (прямой текст), .docx (python-docx), .pdf (pypdf),
.xlsx (openpyxl — РУП часто оформляют таблицей в Excel).
"""
from __future__ import annotations

import io
import os

import docx
from openpyxl import load_workbook
from pypdf import PdfReader

ALLOWED_PLAN_EXT = {".txt", ".docx", ".pdf", ".xlsx"}


def _from_txt(data: bytes) -> str:
    for enc in ("utf-8", "cp1251"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _from_docx(data: bytes) -> str:
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(p.strip() for p in pages if p.strip())


def _from_xlsx(data: bytes) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text(filename: str, data: bytes) -> str:
    """Извлекает текст из файла по его расширению.

    Бросает ValueError для неподдерживаемого формата.
    """
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".txt":
        return _from_txt(data)
    if ext == ".docx":
        return _from_docx(data)
    if ext == ".pdf":
        return _from_pdf(data)
    if ext == ".xlsx":
        return _from_xlsx(data)
    raise ValueError(
        f"Неподдерживаемый формат: {ext}. Разрешено: {sorted(ALLOWED_PLAN_EXT)}"
    )
