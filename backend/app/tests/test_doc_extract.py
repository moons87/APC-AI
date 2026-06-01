import io

import pytest

from app.doc_extract import ALLOWED_PLAN_EXT, extract_text


def test_txt_decoded():
    data = "Сабақ жоспары\nБөлім 1".encode("utf-8")
    assert "Сабақ жоспары" in extract_text("plan.txt", data)


def test_docx_paragraphs_and_tables():
    import docx

    doc = docx.Document()
    doc.add_paragraph("Кіріспе абзац")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Тема"
    table.rows[0].cells[1].text = "Сварка"
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text("plan.docx", buf.getvalue())
    assert "Кіріспе абзац" in text
    assert "Сварка" in text


def test_xlsx_cells_joined():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["№", "Тақырып"])
    ws.append([1, "Дәнекерлеу негіздері"])
    buf = io.BytesIO()
    wb.save(buf)

    text = extract_text("ruP.xlsx", buf.getvalue())
    assert "Тақырып" in text
    assert "Дәнекерлеу негіздері" in text


def test_pdf_pages_joined(monkeypatch):
    class FakePage:
        def __init__(self, t):
            self._t = t

        def extract_text(self):
            return self._t

    class FakeReader:
        def __init__(self, _stream):
            self.pages = [FakePage("Бет 1 мәтіні"), FakePage("Бет 2 мәтіні")]

    monkeypatch.setattr("app.doc_extract.PdfReader", FakeReader)
    text = extract_text("plan.pdf", b"%PDF-fake")
    assert "Бет 1 мәтіні" in text
    assert "Бет 2 мәтіні" in text


def test_unsupported_ext_raises():
    with pytest.raises(ValueError):
        extract_text("plan.rtf", b"data")


def test_allowed_ext_set():
    assert ALLOWED_PLAN_EXT == {".txt", ".docx", ".pdf", ".xlsx"}
